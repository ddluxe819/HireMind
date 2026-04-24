import io
import re
import html as _html


def _e(s) -> str:
    return _html.escape(str(s) if s is not None else "")


def _strip_markers(raw: str) -> str:
    return re.sub(r'\[\[(.+?)\]\]', r'\1', raw)


# ── Plain-text renderer (used for the text download option) ─────────────────

def render_resume_text(d: dict, job_title: str, company: str) -> str:
    lines = []
    name = d.get("name", "")
    lines += [name, d.get("tagline", ""), ""]
    lines += [
        d.get("location", ""), d.get("phone", ""),
        d.get("email", ""), d.get("linkedin", ""), "",
    ]
    lines += ["CORE COMPETENCIES"]
    for c in (d.get("competencies") or []):
        lines.append(f"  {c}")
    lines.append("")
    tech = " · ".join(d.get("tech_stack") or [])
    lines += ["MARTECH STACK", f"  {tech}", ""]
    ai = " · ".join(d.get("ai_tools") or [])
    lines += ["AI & AUTOMATION", f"  {ai}", ""]
    lines += ["THE WORK I DO", d.get("intro_paragraph", ""), ""]
    for item in (d.get("framework_items") or []):
        lines.append(item.get("title", ""))
        lines.append(f"  {item.get('body', '')}")
    lines.append("")
    lines.append("SELECTED ACHIEVEMENTS")
    for grp in (d.get("achievement_groups") or []):
        lines.append(grp.get("title", "").upper())
        for item in (grp.get("items") or []):
            lines.append(f"  – {_strip_markers(item)}")
        lines.append("")
    lines.append("EXPERIENCE")
    for exp in (d.get("experience") or []):
        lines.append(f"{exp.get('title', '')}  {exp.get('dates', '')}")
        lines.append(f"  {exp.get('company', '')}")
    lines.append("")
    edu = d.get("education", "").replace("\n", ", ")
    lines.append(f"EDUCATION: {edu}")
    for cert in (d.get("certifications") or []):
        lines.append(f"  {cert}")
    if d.get("board"):
        lines.append(f"  {d.get('board', '').replace(chr(10), ', ')}")
    header = f"Resume — {name}\nTailored for: {job_title} at {company}\n{'=' * 60}\n"
    return header + "\n".join(lines)


# ── DOCX generator ───────────────────────────────────────────────────────────

def generate_docx_from_data(d: dict, job_title: str, company: str) -> bytes:
    """
    Generate a clean, professional DOCX from resume JSON data.
    Used for tailored resume output when the original is a PDF,
    or as a fallback for DOCX uploads.
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    TEAL = RGBColor(0x2A, 0x7F, 0x7F)
    DARK = RGBColor(0x11, 0x13, 0x1F)
    MID = RGBColor(0x44, 0x44, 0x55)
    LIGHT = RGBColor(0x88, 0x88, 0x99)

    doc = Document()

    # Page margins
    sec = doc.sections[0]
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)
    sec.top_margin = Inches(0.65)
    sec.bottom_margin = Inches(0.6)

    # Remove default paragraph spacing from Normal style
    doc.styles['Normal'].paragraph_format.space_after = Pt(0)
    doc.styles['Normal'].paragraph_format.space_before = Pt(0)

    def _para(text="", bold=False, italic=False, size=10, color=None,
               space_before=0, space_after=3, align=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        if align:
            p.alignment = align
        if text:
            r = p.add_run(text)
            r.bold = bold
            r.italic = italic
            r.font.size = Pt(size)
            if color:
                r.font.color.rgb = color
        return p

    def _section_rule(text):
        """Section header with a teal bottom border."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(text.upper())
        r.bold = True
        r.font.size = Pt(7.5)
        r.font.color.rgb = TEAL
        # Add bottom border via XML
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '2')
        bottom.set(qn('w:color'), '2A7F7F')
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    # ── Name & contact ──
    _para(d.get("name", ""), bold=True, size=22, color=DARK, space_after=2)
    _para(d.get("tagline", ""), bold=True, size=9, color=TEAL, space_after=3)
    contact = "  ·  ".join(
        x for x in [d.get("location"), d.get("phone"), d.get("email"), d.get("linkedin")] if x
    )
    _para(contact, size=9, color=LIGHT, space_after=8)

    # ── The Work I Do ──
    _section_rule("The Work I Do")
    _para(d.get("intro_paragraph", ""), size=10.5, color=MID, space_after=6)

    for item in (d.get("framework_items") or []):
        num = item.get("num", "")
        title = item.get("title", "")
        body = item.get("body", "")
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.1)
        num_r = p.add_run(f"{num}  ")
        num_r.font.size = Pt(7.5)
        num_r.font.color.rgb = TEAL
        title_r = p.add_run(f"{title}  ")
        title_r.bold = True
        title_r.font.size = Pt(9.5)
        title_r.font.color.rgb = DARK
        body_r = p.add_run(body)
        body_r.font.size = Pt(9)
        body_r.font.color.rgb = MID

    # ── Selected Achievements ──
    _section_rule("Selected Achievements")
    for grp in (d.get("achievement_groups") or []):
        _para(grp.get("title", "").upper(), bold=True, size=8, color=TEAL,
              space_before=6, space_after=3)
        for item_text in (grp.get("items") or []):
            clean = _strip_markers(item_text)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Inches(0.2)
            r = p.add_run(clean)
            r.font.size = Pt(10)
            r.font.color.rgb = MID

    # ── Experience ──
    _section_rule("Experience")
    for exp in (d.get("experience") or []):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(1)
        title_r = p.add_run(exp.get("title", ""))
        title_r.bold = True
        title_r.font.size = Pt(11)
        title_r.font.color.rgb = DARK
        dates_r = p.add_run(f"   {exp.get('dates', '')}")
        dates_r.font.size = Pt(9)
        dates_r.font.color.rgb = LIGHT
        co_p = doc.add_paragraph()
        co_p.paragraph_format.space_after = Pt(1)
        co_r = co_p.add_run(exp.get("company", ""))
        co_r.font.size = Pt(9.5)
        co_r.font.color.rgb = TEAL

    # ── Education ──
    _section_rule("Education")
    edu = (d.get("education") or "").replace("\n", "   ")
    _para(edu, size=10, color=DARK, space_before=2)
    for cert in (d.get("certifications") or []):
        _para(cert, size=9, color=MID)
    if d.get("board"):
        _para(d.get("board", "").replace("\n", "   "), size=9, color=LIGHT)

    # ── Skills & Tools ──
    _section_rule("Skills & Tools")
    comps = "  ·  ".join(d.get("competencies") or [])
    if comps:
        _para(comps, size=9.5, color=MID, space_before=2)
    tech = "  ·  ".join(d.get("tech_stack") or [])
    if tech:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        lbl = p.add_run("Martech:  ")
        lbl.bold = True
        lbl.font.size = Pt(8.5)
        lbl.font.color.rgb = TEAL
        val = p.add_run(tech)
        val.font.size = Pt(8.5)
        val.font.color.rgb = MID
    ai = "  ·  ".join(d.get("ai_tools") or [])
    if ai:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        lbl = p.add_run("AI & Automation:  ")
        lbl.bold = True
        lbl.font.size = Pt(8.5)
        lbl.font.color.rgb = TEAL
        val = p.add_run(ai)
        val.font.size = Pt(8.5)
        val.font.color.rgb = MID

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── DOCX surgical editor ──────────────────────────────────────────────────────

def edit_docx_with_replacements(docx_bytes: bytes, replacements: list) -> bytes:
    """
    Apply find/replace pairs to an existing DOCX, preserving paragraph styles.
    Consolidates runs within each paragraph when a match is found.
    """
    from docx import Document

    doc = Document(io.BytesIO(docx_bytes))

    def _apply(para, find, replace):
        full = "".join(r.text for r in para.runs)
        if find not in full:
            return
        new_text = full.replace(find, replace, 1)
        if para.runs:
            para.runs[0].text = new_text
            for r in para.runs[1:]:
                r.text = ""

    for repl in (replacements or []):
        find = repl.get("find", "")
        replace = repl.get("replace", "")
        if not find:
            continue
        for para in doc.paragraphs:
            _apply(para, find, replace)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _apply(para, find, replace)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
